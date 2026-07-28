from __future__ import annotations

import unittest
from io import BytesIO
from types import SimpleNamespace
from unittest.mock import MagicMock, patch

from workers import tasks


class DocumentWorkerTaskTests(unittest.TestCase):
    def test_plain_text_extraction_returns_one_chunk(self) -> None:
        self.assertEqual(
            tasks._extract_chunks("text/plain", "第一段\n第二段".encode()),
            [("第一段\n第二段", None)],
        )

    def test_pdf_extraction_preserves_page_numbers(self) -> None:
        import fitz

        document = fitz.open()
        page = document.new_page()
        page.insert_text((72, 72), "PDF content")

        self.assertEqual(tasks._extract_chunks("application/pdf", document.tobytes()), [("PDF content", 1)])

    def test_docx_extraction_collects_paragraphs(self) -> None:
        from docx import Document

        document = Document()
        document.add_paragraph("第一段")
        document.add_paragraph("第二段")
        output = BytesIO()
        document.save(output)

        self.assertEqual(
            tasks._extract_chunks(
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document", output.getvalue(),
            ),
            [("第一段\n第二段", None)],
        )

    @patch("workers.tasks.knowledge_repository.complete_ingestion")
    @patch("workers.tasks.read_private_object")
    @patch("workers.tasks.knowledge_repository.mark_ingestion_processing")
    @patch("workers.tasks.knowledge_repository.get_ingestion_job")
    def test_processing_job_marks_document_ready_after_text_extraction(
        self,
        get_job: MagicMock,
        mark_processing: MagicMock,
        read_object: MagicMock,
        complete_ingestion: MagicMock,
    ) -> None:
        get_job.return_value = {"storage_key": "knowledge-files/user-1/file.txt", "mime_type": "text/plain"}
        response = MagicMock()
        response.read.return_value = "可解析内容".encode()
        read_object.return_value = response

        tasks.process_document_ingestion("job-1")

        mark_processing.assert_called_once_with("job-1")
        complete_ingestion.assert_called_once_with("job-1", [("可解析内容", None)])
        response.close.assert_called_once()
        response.release_conn.assert_called_once()

    @patch("workers.tasks.remove_private_object")
    @patch("workers.tasks.knowledge_repository.delete_node_tree", return_value=["knowledge-files/user-1/file.txt"])
    @patch("workers.tasks.knowledge_repository.fail_ingestion")
    @patch("workers.tasks.read_private_object")
    @patch("workers.tasks.knowledge_repository.mark_ingestion_processing")
    @patch("workers.tasks.knowledge_repository.get_ingestion_job")
    def test_failed_processing_removes_file_from_knowledge_base(
        self,
        get_job: MagicMock,
        _: MagicMock,
        read_object: MagicMock,
        fail_ingestion: MagicMock,
        delete_node_tree: MagicMock,
        remove_object: MagicMock,
    ) -> None:
        get_job.return_value = {
            "job_id": "job-1", "node_id": "node-1", "owner_user_id": "user-1",
            "storage_key": "knowledge-files/user-1/file.txt", "mime_type": "text/plain",
        }
        response = MagicMock()
        response.read.return_value = b"\n"
        read_object.return_value = response

        with self.assertRaisesRegex(ValueError, "no extractable text"):
            tasks.process_document_ingestion("job-1")

        fail_ingestion.assert_called_once()
        delete_node_tree.assert_called_once_with("node-1", "user-1")
        remove_object.assert_called_once_with("knowledge-files/user-1/file.txt")

    @patch("workers.tasks.system", return_value="Linux")
    @patch("workers.tasks.get_object_storage_settings")
    @patch("workers.tasks.create_object_storage_client")
    @patch("workers.tasks.get_connection")
    def test_probe_reports_database_and_private_bucket_reachability(
        self,
        get_connection: MagicMock,
        create_client: MagicMock,
        get_settings: MagicMock,
        _: MagicMock,
    ) -> None:
        connection = MagicMock()
        get_connection.return_value.__enter__.return_value = connection
        get_settings.return_value = SimpleNamespace(bucket="ai-platform-private")
        create_client.return_value.bucket_exists.return_value = True

        result = tasks.run_document_processing_probe()

        self.assertEqual(result["database"], "ok")
        self.assertEqual(result["object_storage_bucket"], "ai-platform-private")
        connection.cursor.return_value.__enter__.return_value.execute.assert_called_once_with("SELECT 1")
        create_client.return_value.bucket_exists.assert_called_once_with("ai-platform-private")

    @patch("workers.tasks.get_object_storage_settings")
    @patch("workers.tasks.create_object_storage_client")
    @patch("workers.tasks.get_connection")
    def test_probe_rejects_missing_private_bucket(
        self,
        get_connection: MagicMock,
        create_client: MagicMock,
        get_settings: MagicMock,
    ) -> None:
        get_settings.return_value = SimpleNamespace(bucket="ai-platform-private")
        create_client.return_value.bucket_exists.return_value = False

        with self.assertRaisesRegex(RuntimeError, "cannot find private bucket"):
            tasks.run_document_processing_probe()

        self.assertTrue(get_connection.called)


if __name__ == "__main__":
    unittest.main()
