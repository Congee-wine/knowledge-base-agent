from __future__ import annotations

import unittest
from io import BytesIO
from unittest.mock import MagicMock, patch

from docx import Document

from services import document_preview
from services.errors import DomainError


def _document(mime_type: str, status: str = "ready") -> dict[str, str]:
    return {
        "storage_key": "knowledge-files/user-1/document",
        "mime_type": mime_type,
        "name": "document.docx",
        "processing_status": status,
    }


class DocumentPreviewTests(unittest.TestCase):
    @patch("services.document_preview.read_private_object")
    @patch("services.document_preview.knowledge_repository.find_owned_current_file_version")
    def test_text_preview_reads_only_ready_owned_document(self, find_document: MagicMock, read_object: MagicMock) -> None:
        find_document.return_value = _document("text/markdown")
        response = MagicMock()
        response.read.return_value = b"# title"
        read_object.return_value = response

        preview = document_preview.get_document_preview("user-1", "node-1")

        self.assertEqual(preview.kind, "text")
        self.assertEqual(preview.content, "# title")
        self.assertTrue(preview.is_markdown)
        find_document.assert_called_once_with("node-1", "user-1")
        response.close.assert_called_once()
        response.release_conn.assert_called_once()

    @patch("services.document_preview.knowledge_repository.find_owned_current_file_version")
    def test_preview_rejects_unready_document(self, find_document: MagicMock) -> None:
        find_document.return_value = _document("text/plain", "processing")

        with self.assertRaises(DomainError) as captured:
            document_preview.get_document_preview("user-1", "node-1")

        self.assertEqual(captured.exception.code, "DOCUMENT_NOT_READY")

    def test_docx_html_escapes_text_and_keeps_safe_structure(self) -> None:
        document = Document()
        document.add_heading("<script>alert(1)</script>", level=1)
        document.add_paragraph("plain <b>text</b>")
        table = document.add_table(rows=1, cols=1)
        table.cell(0, 0).text = "<img src=x>"
        output = BytesIO()
        document.save(output)

        html = document_preview._convert_docx_to_safe_html(output.getvalue())

        self.assertIn("&lt;script&gt;alert(1)&lt;/script&gt;", html)
        self.assertIn("&lt;b&gt;text&lt;/b&gt;", html)
        self.assertIn("&lt;img src=x&gt;", html)
        self.assertNotIn("<script>", html)
        self.assertIn("<table>", html)

    @patch("services.document_preview.read_private_object", side_effect=RuntimeError("MinIO unavailable"))
    @patch("services.document_preview.knowledge_repository.find_owned_current_file_version")
    def test_storage_failure_returns_preview_unavailable(self, find_document: MagicMock, _: MagicMock) -> None:
        find_document.return_value = _document("application/pdf")

        with self.assertRaises(DomainError) as captured:
            document_preview.get_document_preview("user-1", "node-1")

        self.assertEqual(captured.exception.code, "DOCUMENT_PREVIEW_UNAVAILABLE")


if __name__ == "__main__":
    unittest.main()
