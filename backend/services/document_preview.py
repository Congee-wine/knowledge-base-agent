from __future__ import annotations

from dataclasses import dataclass
from html import escape
from io import BytesIO
from typing import Literal

from integrations.object_storage import read_private_object
from repositories import knowledge as knowledge_repository
from services.errors import (
    document_not_ready,
    document_preview_empty,
    document_preview_unavailable,
    not_found,
    unsupported_document_type,
)


PreviewKind = Literal["pdf", "text", "html"]


@dataclass(frozen=True)
class DocumentPreview:
    kind: PreviewKind
    name: str
    content: bytes | str
    is_markdown: bool = False


def get_document_preview(user_id: str, node_id: str) -> DocumentPreview:
    document = knowledge_repository.find_owned_current_file_version(node_id, user_id)
    if document is None:
        raise not_found()
    if document["processing_status"] != "ready":
        raise document_not_ready()
    content = _read_document_content(str(document["storage_key"]))
    mime_type = str(document["mime_type"])
    name = str(document["name"])
    if mime_type == "application/pdf":
        return DocumentPreview("pdf", name, content)
    if "wordprocessingml" in mime_type:
        return DocumentPreview("html", name, _convert_docx_to_safe_html(content))
    if mime_type in {"text/plain", "text/markdown"}:
        text = _decode_text(content)
        return DocumentPreview("text", name, text, mime_type == "text/markdown" or name.lower().endswith((".md", ".markdown")))
    raise unsupported_document_type()


def _read_document_content(storage_key: str) -> bytes:
    try:
        response = read_private_object(storage_key)
        try:
            return response.read()
        finally:
            response.close()
            response.release_conn()
    except Exception as error:
        raise document_preview_unavailable() from error


def _decode_text(content: bytes) -> str:
    try:
        text = content.decode("utf-8").strip()
    except UnicodeDecodeError as error:
        raise document_preview_empty() from error
    if not text:
        raise document_preview_empty()
    return text


def _convert_docx_to_safe_html(content: bytes) -> str:
    try:
        from docx import Document

        document = Document(BytesIO(content))
        fragments = _render_paragraphs(document.paragraphs)
        fragments.extend(_render_table(table) for table in document.tables)
    except Exception as error:
        raise document_preview_unavailable() from error
    html = "".join(fragment for fragment in fragments if fragment)
    if not html:
        raise document_preview_empty()
    return html


def _render_paragraphs(paragraphs: object) -> list[str]:
    fragments: list[str] = []
    list_items: list[str] = []
    list_tag: str | None = None
    for paragraph in paragraphs:
        style_name = str(paragraph.style.name).lower()
        text = escape(paragraph.text.strip())
        is_list = style_name.startswith("list")
        if is_list and text:
            tag = "ol" if "number" in style_name else "ul"
            if list_tag and list_tag != tag:
                fragments.append(f"<{list_tag}>{''.join(list_items)}</{list_tag}>")
                list_items = []
            list_tag = tag
            list_items.append(f"<li>{text}</li>")
            continue
        if list_tag:
            fragments.append(f"<{list_tag}>{''.join(list_items)}</{list_tag}>")
            list_items, list_tag = [], None
        if text:
            fragments.append(_render_paragraph(style_name, text))
    if list_tag:
        fragments.append(f"<{list_tag}>{''.join(list_items)}</{list_tag}>")
    return fragments


def _render_paragraph(style_name: str, text: str) -> str:
    if style_name.startswith("heading"):
        level = next((character for character in style_name if character.isdigit()), "1")
        return f"<h{level}>{text}</h{level}>"
    return f"<p>{text}</p>"


def _render_table(table: object) -> str:
    rows = []
    for row in table.rows:
        cells = "".join(f"<td>{escape(cell.text.strip())}</td>" for cell in row.cells)
        rows.append(f"<tr>{cells}</tr>")
    return f"<table><tbody>{''.join(rows)}</tbody></table>" if rows else ""
